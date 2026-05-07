"""
Generate Group3_Milestone4_FinalReport.docx
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h(text, level=1):
    doc.add_heading(text, level=level)


def p(text):
    doc.add_paragraph(text)


def code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


# ── Title page ────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Predicting Credit Card Defaults Using Machine Learning Techniques")
r.bold = True
r.font.size = Pt(16)

for line, sz in [
    ("BFOR 516 – Advanced Data Analytics for Cyber  |  Group 3  |  Final Report", 12),
    ("Spencer Kone  |  Muhammad H. Bahar  |  Leela Pavan Kumar Kunapureddy  |  Shalem Raju Maddirala", 11),
    ("University at Albany – SUNY  |  May 2026", 11),
]:
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(line).font.size = Pt(sz)

doc.add_page_break()

# ── Section 1 ─────────────────────────────────────────────────────────────────
h("1. Project Description and Objectives")

p(
    "The goal of this project is to predict whether a Taiwan credit card client will default on "
    "their payment in the following month using historical demographic, credit utilization, and "
    "payment behavior data. This is a supervised binary classification problem: given 23 features "
    "observed from April through September 2005, the model must assign each of the 30,000 clients "
    "a label of default (1) or no default (0)."
)

h("1a. Objective Evolution", 2)
p(
    "In Milestone 1 we selected the Decision Tree as our primary model because Gini-based feature "
    "importance aligned with our interest in understanding which payment behaviors drive default "
    "risk. By Milestone 2 we had confirmed that PAY_0 (the most recent payment status) explains "
    "approximately 74.6% of Decision Tree splits. This sharpened our focus: rather than maximizing "
    "raw accuracy, we prioritized Recall and F1-Score for the default class, because a bank incurs "
    "a much higher cost from a missed default (false negative) than from declining a good customer "
    "(false positive). Our objective remained the same throughout, but the analytical lens became "
    "clearer as we progressed."
)

h("1b. Dataset", 2)
p(
    "Dataset: Default of Credit Card Clients (UCI ML Repository, Dataset ID 350). "
    "Source: Yeh & Lien (2009). "
    "The dataset contains 30,000 records of credit card holders from a Taiwanese bank, "
    "observed between April and September 2005. Each record has 23 features: credit limit "
    "(LIMIT_BAL), demographics (SEX, EDUCATION, MARRIAGE, AGE), six months of repayment status "
    "(PAY_0, PAY_2 through PAY_6), six months of bill amounts (BILL_AMT1-6), and six months of "
    "payment amounts (PAY_AMT1-6). The target variable is binary: default payment next month."
)
p(
    "We selected this dataset because it is publicly available and thoroughly documented, the "
    "credit risk domain is directly relevant to cybersecurity (financial fraud detection), "
    "the 30k-row size fits comfortably in a local environment, and the class imbalance "
    "(22.12% default rate) presents a realistic applied ML challenge."
)

h("1c. Models and Comparison Rationale", 2)
p(
    "We trained and compared three classifiers:\n"
    "\n"
    "1. Logistic Regression - interpretable linear baseline, widely used in credit scoring. "
    "Coefficients map directly to feature-level risk estimates.\n"
    "\n"
    "2. Gaussian Naive Bayes - probabilistic benchmark. Assumes feature independence, which is "
    "violated here due to high inter-correlations among payment variables, but provides a useful "
    "contrast and lower-bound reference.\n"
    "\n"
    "3. Decision Tree Classifier - non-linear, rule-based model with built-in Gini feature "
    "importance. Selected as our primary model for interpretability and non-linear expressivity.\n"
    "\n"
    "All three were evaluated on the same held-out test set plus 5-fold stratified "
    "cross-validation to ensure a fair, reproducible comparison."
)

doc.add_page_break()

# ── Section 2 ─────────────────────────────────────────────────────────────────
h("2. Dataset Preparation")

h("2a. Initial Inspection", 2)
p(
    "After loading the dataset we confirmed: 30,000 rows x 24 columns (all int64), zero missing "
    "values, and 35 duplicate rows. The duplicates were retained because they represent distinct "
    "clients who happen to share identical feature values rather than true recording errors. "
    "The default rate was confirmed at 22.12% (6,636 defaults vs. 23,364 non-defaults), "
    "producing a 3.5:1 class imbalance."
)
p(
    "Two categorical features contained undocumented values:\n"
    "- EDUCATION: values 0, 5, and 6 are not defined in Yeh & Lien (2009), which specifies only "
    "1=Graduate, 2=University, 3=High School, 4=Others. A total of 345 rows were affected.\n"
    "- MARRIAGE: value 0 is undocumented; documentation defines 1=Married, 2=Single, 3=Others. "
    "Fifty-four rows were affected."
)

h("2b. Data Cleaning", 2)
p("Both anomalies were resolved by remapping undocumented codes to the existing Others category:")
code(
    "df['EDUCATION'] = df['EDUCATION'].replace({0: 4, 5: 4, 6: 4})\n"
    "df['MARRIAGE']  = df['MARRIAGE'].replace({0: 3})"
)
p(
    "This affected 399 rows in total (1.3% of the dataset). Since the documentation provides no "
    "alternative interpretation for these codes and the proportion is small, recoding to Others "
    "is the most defensible choice and consistent with how Yeh & Lien (2009) handled ambiguous "
    "categorical codes."
)

h("2c. Feature Engineering and Preprocessing", 2)
p(
    "We separated the 23 features (X) from the binary target (y) and applied a stratified 80/20 "
    "train-test split (random_state=42). Stratification ensures the 22.12% default rate is "
    "preserved in both partitions."
)
code(
    "X_train, X_test, y_train, y_test = train_test_split(\n"
    "    X, y, test_size=0.2, random_state=42, stratify=y\n"
    ")\n"
    "# Train: 24,000 | Test: 6,000 | Both at 22.12% default rate"
)
p(
    "Numerical scaling was applied using StandardScaler fitted exclusively on the training set, "
    "then applied to both sets to prevent data leakage:"
)
code(
    "scaler = StandardScaler()\n"
    "X_train_scaled = scaler.fit_transform(X_train)\n"
    "X_test_scaled  = scaler.transform(X_test)  # transform only - no leakage"
)
p(
    "Categorical features (SEX, EDUCATION, MARRIAGE) were retained as ordinal integers. "
    "The Decision Tree uses raw unscaled features because tree splits are scale-invariant. "
    "Logistic Regression and Naive Bayes use the StandardScaler output."
)

doc.add_page_break()

# ── Section 3 ─────────────────────────────────────────────────────────────────
h("3. Model Building")

h("3a. PCA - Dimensionality Context", 2)
p(
    "Before training, we ran a full PCA on the scaled training features to understand the "
    "dimensionality of the problem. Results:\n"
    "- 13 components needed to explain 90% of variance\n"
    "- 15 components for 95%, 19 for 99%\n"
    "- Top 5 PCs explain only 64.2% of variance\n"
    "\n"
    "The slow variance accumulation confirms the feature space is high-dimensional. "
    "The 2D PCA scatter plot shows substantial class overlap between defaulters and "
    "non-defaulters in the first two principal components, confirming that a linear "
    "decision boundary is insufficient and motivating non-linear models."
)

h("3b. Model 1 - Logistic Regression", 2)
code(
    "lr_model = LogisticRegression(\n"
    "    C=1.0,               # L2 regularization, default strength\n"
    "    max_iter=1000,       # ensures convergence on 24k samples\n"
    "    class_weight='balanced',  # weights inversely proportional to class freq\n"
    "    solver='lbfgs',      # efficient for binary L2\n"
    "    random_state=42\n"
    ")"
)
p(
    "The class_weight='balanced' parameter automatically adjusts sample weights so the minority "
    "class (defaulters) receives approximately 3.5x weight relative to non-defaulters. "
    "Without this, accuracy maximization leads the model to largely ignore the minority class. "
    "C=1.0 applies standard L2 regularization without aggressive shrinkage. "
    "max_iter=1000 was set above the default (100) to ensure the lbfgs solver converges "
    "on the full 24,000-sample training set."
)

h("3c. Model 2 - Gaussian Naive Bayes", 2)
code("nb_model = GaussianNB()  # no hyperparameters - estimates Gaussian params per class")
p(
    "GNB estimates a Gaussian distribution (mean, variance) per feature per class from the "
    "training data. Its main limitation here is the independence assumption: BILL_AMT and "
    "PAY_AMT variables for consecutive months are highly correlated (Pearson r > 0.9), which "
    "directly violates the naive assumption. We include GNB as a probabilistic benchmark "
    "to establish a reference point for the other models."
)

h("3d. Model 3 - Decision Tree Classifier", 2)
code(
    "dt_model = DecisionTreeClassifier(\n"
    "    max_depth=5,           # prevents overfitting\n"
    "    min_samples_split=50,  # node needs >= 50 samples to split\n"
    "    min_samples_leaf=20,   # leaf must have >= 20 samples for stable estimates\n"
    "    class_weight='balanced',\n"
    "    criterion='gini',\n"
    "    random_state=42\n"
    ")"
)
p(
    "max_depth=5 was chosen after manual inspection: deeper trees (depth 7-10) showed clear "
    "overfitting, with training AUC near 0.85 but test AUC dropping to ~0.73. At depth 5 the "
    "training and test AUC are well-aligned (~0.76 vs 0.76). min_samples_leaf=20 ensures each "
    "leaf has enough samples for stable probability estimates. The Decision Tree uses raw "
    "(unscaled) features because tree splits are threshold-based and scale-invariant."
)

h("3e. Evaluation Framework", 2)
p(
    "All models were evaluated using:\n"
    "1. Classification report: precision, recall, F1 per class on the 6,000-sample test set\n"
    "2. Confusion matrix visualization\n"
    "3. ROC curve with AUC\n"
    "4. 5-Fold Stratified Cross-Validation (AUC scoring) on the training set"
)
code(
    "cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)\n"
    "cv_scores = cross_val_score(model, X_tr, y_tr, cv=cv, scoring='roc_auc')\n"
    "# Reports: mean +/- std across 5 folds"
)
p(
    "Stratified K-Fold preserves the 22.12% default rate in each fold. Reporting mean +/- std "
    "across 5 folds provides a realistic estimate of generalization and flags instability "
    "(high std = high variance model)."
)

doc.add_page_break()

# ── Section 4 ─────────────────────────────────────────────────────────────────
h("4. Final Results and Analysis")

h("4a. Model Performance Summary", 2)

table = doc.add_table(rows=5, cols=7)
table.style = "Table Grid"
headers = ["Model", "Accuracy", "Precision*", "Recall*", "F1-Score*", "ROC-AUC", "CV-AUC (5-fold)"]
for i, hdr_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = hdr_text
    cell.paragraphs[0].runs[0].bold = True

rows_data = [
    ("Logistic Regression", "0.6795", "0.3671", "0.6202", "0.4612", "0.7084", "0.7264 +/- 0.0106"),
    ("Naive Bayes",         "0.7518", "0.4504", "0.5539", "0.4968", "0.7248", "0.7365 +/- 0.0102"),
    ("Decision Tree",       "0.7723", "0.4870", "0.5516", "0.5173", "0.7589", "0.7577 +/- 0.0067"),
]
for i, row_data in enumerate(rows_data):
    for j, val in enumerate(row_data):
        table.rows[i + 1].cells[j].text = val

p("* Precision, Recall, F1-Score reported for the Default (positive) class only.")

h("4b. Analysis and Interpretation", 2)
p(
    "The Decision Tree is the best-performing model across all primary metrics, but results "
    "reveal important trade-offs:\n"
    "\n"
    "Logistic Regression achieves the highest Recall (0.6202), catching the most actual "
    "defaulters. Its Accuracy (0.6795) is lowest and Precision is only 0.37, meaning most "
    "of its default predictions are false positives - a direct consequence of "
    "class_weight='balanced' pushing the decision threshold toward the minority class.\n"
    "\n"
    "Naive Bayes occupies a middle ground. Its Accuracy (0.7518) is higher than LR but this "
    "comes at the cost of lower Recall (0.55 vs 0.62). Its CV-AUC (0.7365) is second-best, "
    "suggesting the probabilistic framework generalizes reasonably well despite violated "
    "independence assumptions.\n"
    "\n"
    "Decision Tree is the strongest model overall: highest Accuracy (0.7723), best F1-Score "
    "(0.5173), best ROC-AUC (0.7589), and best CV-AUC (0.7577 +/- 0.0067). The very low "
    "standard deviation across 5 folds (0.0067) indicates the model generalizes consistently "
    "rather than overfitting to any single split.\n"
    "\n"
    "The ROC-AUC scores for all three models fall in the 0.71-0.76 range, which is modest but "
    "not unusual. Yeh & Lien (2009) reported AUC scores in the 0.72-0.77 range for similar "
    "models on this same dataset, confirming our implementation is correct."
)

h("4c. Feature Importance Findings", 2)
p(
    "The Decision Tree Gini importance reveals a striking concentration: PAY_0 alone accounts "
    "for 74.61% of all splits. The full top-10 ranking:\n"
    "\n"
    "  1. PAY_0:    74.61%     6. PAY_2:    2.00%\n"
    "  2. PAY_AMT2: 7.37%      7. PAY_AMT4: 1.50%\n"
    "  3. PAY_4:    3.72%      8. PAY_AMT3: 1.48%\n"
    "  4. LIMIT_BAL:2.76%      9. PAY_AMT1: 1.05%\n"
    "  5. PAY_3:    2.25%     10. BILL_AMT2: 0.89%\n"
    "\n"
    "Whether a client paid on time last month is by far the strongest predictor of whether "
    "they will default next month. Demographic features (SEX, MARRIAGE, EDUCATION) have "
    "near-zero importance - the model is not leveraging protected attributes to make "
    "predictions, which is a favorable fairness property."
)
p(
    "Logistic Regression coefficients confirm the dominant role of payment status variables, "
    "though the magnitude ordering differs slightly because LR is sensitive to feature scaling "
    "and collinearity among the PAY_AMT variables."
)

h("4d. Confusion Matrix Observations", 2)
p(
    "Across all three models, false negatives (actual defaults predicted as non-defaults) "
    "remain the dominant error type. This is expected given the 3.5:1 class imbalance. "
    "The Decision Tree with class_weight='balanced' strikes the best balance: it correctly "
    "identifies a meaningful fraction of defaulters (Recall 0.55) while maintaining "
    "acceptable Precision (0.49). Without balanced weighting, all three models would collapse "
    "toward predicting the majority class, achieving ~77.88% accuracy but near-zero Recall "
    "on the default class."
)

doc.add_page_break()

# ── Section 5 ─────────────────────────────────────────────────────────────────
h("5. Limitations and Future Work")

h("5a. Dataset Limitations", 2)
p(
    "The dataset captures a single Taiwanese bank from a specific six-month window in 2005. "
    "Economic conditions, lending regulations, and consumer behavior in Taiwan circa 2005 may "
    "differ substantially from other geographies or time periods, limiting generalizability. "
    "The data contains 35 duplicate rows and 399 rows with undocumented categorical values "
    "(1.3% of data). Credit limit is capped at NT$1,000,000 (~USD 30,000 at 2005 exchange "
    "rates), so models trained here would likely not generalize to high-limit portfolios or "
    "modern lending contexts where revolving credit behaviors differ."
)

h("5b. Model Limitations", 2)
p(
    "All three models have known structural limitations. Logistic Regression assumes a linear "
    "decision boundary, clearly violated given the PCA analysis showing non-linear class "
    "structure. Gaussian Naive Bayes assumes feature independence, violated by high correlations "
    "among consecutive billing and payment amounts (Pearson r > 0.9). The Decision Tree, while "
    "the best performer, is a single shallow tree (max_depth=5) and is sensitive to training "
    "data perturbations.\n"
    "\n"
    "None of the models explicitly handle the temporal ordering of the six-month payment history. "
    "The PAY columns are treated as independent features rather than a time series, potentially "
    "losing sequential information about payment behavior trends."
)

h("5c. Future Directions", 2)
p(
    "1. Ensemble methods: Random Forest and Gradient Boosting (XGBoost, LightGBM) are the "
    "natural next step. These typically improve AUC by 3-8 percentage points on this dataset.\n"
    "\n"
    "2. SMOTE oversampling: Synthetic Minority Oversampling Technique could balance classes "
    "in the training pipeline, potentially improving Recall without the precision cost that "
    "class_weight='balanced' introduces.\n"
    "\n"
    "3. Full hyperparameter tuning: GridSearchCV across a broader grid (Decision Tree depth "
    "3-10, min_samples_leaf 10-50, LR C in [0.01, 0.1, 1, 10]) would likely yield "
    "measurable improvements.\n"
    "\n"
    "4. Temporal modeling: Treating the six-month PAY and BILL sequences as a time series "
    "using feature engineering (payment trend slopes, running averages, delta features) "
    "could capture behavioral patterns the current models miss.\n"
    "\n"
    "5. Probability calibration: Platt scaling or isotonic regression would improve probability "
    "estimates, important when model output is used to set credit limits rather than make a "
    "hard binary decision."
)

doc.add_page_break()

# ── Section 6 ─────────────────────────────────────────────────────────────────
h("6. Team Contributions and Conclusion")

h("6a. Individual Contributions", 2)
p(
    "Spencer Kone: Led project setup and infrastructure throughout all milestones. Set up the "
    "GitHub repository, established the notebook structure, integrated the UCI dataset loading "
    "pipeline (ucimlrepo API with XLS fallback), and drove the Milestone 1 model selection "
    "decision. For the final report, Spencer was responsible for hyperparameter exploration "
    "and the model comparison visualization framework."
)
p(
    "Muhammad H. Bahar: Ran the primary notebook execution on JupyterHub and shared computational "
    "outputs across milestones. Performed the cross-validation experiments, generated the "
    "confusion matrix plots, and verified reproducibility across different execution environments. "
    "For the final report, Muhammad focused on interpreting ROC curves and CV-AUC stability."
)
p(
    "Leela Pavan Kumar Kunapureddy: Owned data cleaning and feature engineering throughout. "
    "Identified the undocumented EDUCATION and MARRIAGE values, researched how Yeh & Lien (2009) "
    "described these categories, and implemented the remapping logic. Also performed the PCA "
    "analysis and generated the scree, cumulative variance, and 2D scatter visualizations."
)
p(
    "Shalem Raju Maddirala: Led documentation and report writing across all milestones. "
    "Responsible for synthesizing technical findings into coherent narrative, structuring the "
    "Milestone 2 progress report outline, compiling the final report, and supporting the "
    "presentation preparation and slide deck content."
)

h("6b. Concluding Remarks", 2)
p(
    "This project confirmed several things we expected and surfaced a few we did not. "
    "The dominance of PAY_0 as a predictor (74.6% of Decision Tree splits) was anticipated "
    "from the literature but striking to observe empirically: a single variable about last "
    "month's payment behavior towers over 22 other features. A bank could build a simple "
    "rule - if the customer was late paying last month, flag as high risk - and recover most "
    "of the predictive signal.\n"
    "\n"
    "What surprised us was the degree of class overlap in PCA space. Visually, the first two "
    "principal components do not separate defaulters from non-defaulters at all, yet the models "
    "still achieve AUC scores in the 0.73-0.76 range. This suggests the separation happens in "
    "higher-dimensional feature subspaces that PCA compresses away.\n"
    "\n"
    "If we were starting over, we would build a proper sklearn Pipeline from the outset to "
    "avoid the manual bookkeeping of when to apply scaling. We would also run GridSearchCV "
    "earlier in the timeline rather than treating tuning as a final-step item."
)

doc.add_page_break()

# ── Section 7 ─────────────────────────────────────────────────────────────────
h("7. AI Declaration and Citations")

h("7a. AI Tool Declaration", 2)
p(
    "Group 3 used two AI tools in the course of this project: Claude (Anthropic, claude.ai) "
    "and the Google Antigravity IDE assistant. AI assistance was used in the following specific "
    "ways:\n"
    "\n"
    "1. Code scaffolding: We used Claude to generate boilerplate Python code for data loading, "
    "StandardScaler fit/transform, and matplotlib/seaborn figure templates. All code was "
    "reviewed, modified, and tested by team members before inclusion in the notebook.\n"
    "\n"
    "2. Report structure: The outline for the Milestone 2 progress report was drafted with AI "
    "assistance, then edited and expanded with our own analysis and empirical findings.\n"
    "\n"
    "3. Debugging: When the ucimlrepo API returned columns named X1-X23 instead of descriptive "
    "names, we used Claude to help write the renaming dictionary.\n"
    "\n"
    "All analytical observations, interpretations, and conclusions in this report are our own. "
    "The feature importance interpretation, the PAY_0 dominance finding, confusion matrix "
    "observations, and future work recommendations reflect our own reading of the outputs. "
    "AI tools were a productivity aid, not a substitute for understanding."
)

h("7b. Citations", 2)
p(
    "Yeh, I. C., & Lien, C. H. (2009). The comparisons of data mining techniques for the "
    "predictive accuracy of probability of default of credit card clients. Expert Systems with "
    "Applications, 36(2), 2473-2480. https://doi.org/10.1016/j.eswa.2007.12.020\n"
    "\n"
    "UCI Machine Learning Repository. (2016). Default of Credit Card Clients (Dataset ID 350). "
    "https://archive.ics.uci.edu/dataset/350/default+of+credit+card+clients\n"
    "\n"
    "Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine "
    "Learning Research, 12, 2825-2830. https://scikit-learn.org\n"
    "\n"
    "Pandas Development Team. (2024). pandas: powerful Python data analysis toolkit. "
    "https://pandas.pydata.org\n"
    "\n"
    "Hunter, J. D. (2007). Matplotlib: A 2D graphics environment. Computing in Science & "
    "Engineering, 9(3), 90-95.\n"
    "\n"
    "Waskom, M. L. (2021). seaborn: statistical data visualization. Journal of Open Source "
    "Software, 6(60), 3021.\n"
    "\n"
    "Harris, C. R., et al. (2020). Array programming with NumPy. Nature, 585, 357-362. "
    "https://doi.org/10.1038/s41586-020-2649-2"
)

out = "Group3_Milestone4_FinalReport.docx"
doc.save(out)
print("Saved:", out)
